#!/usr/bin/env python3
"""
O'zbekiston OTM Ta'lim Boti v3
Reja → Bo'limlar (oqimli ilmiy matn) → Word fayl
"""

import os, io, re, logging, json
import google.generativeai as genai
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    CallbackQueryHandler, ConversationHandler,
    filters, ContextTypes
)
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(format="%(asctime)s - %(levelname)s - %(message)s", level=logging.INFO)
logger = logging.getLogger(__name__)

BOT_TOKEN = os.environ["BOT_TOKEN"]
GEMINI_KEY = os.environ["GEMINI_API_KEY"]
genai.configure(api_key=GEMINI_KEY)
gemini = genai.GenerativeModel(
    model_name="gemini-2.0-flash-lite",
    generation_config={"max_output_tokens": 8192, "temperature": 0.7}
)

LANG, WORK_TYPE, UNIVERSITY, FACULTY, DIRECTION, SUBJECT, STUDENT_NAME, COURSE, STUDY_TYPE, TOPIC = range(10)

T = {
    "uz": {
        "welcome": "👋 *Assalomu alaykum!*\n\nMen O'zbekiston OTM standartlarida yozma ishlar tayyorlab beruvchi botman.\n\n📝 *Mustaqil ish* — 15–20 sahifa\n📚 *Kurs ishi* — 25–35 sahifa\n\n🌐 *Tilni tanlang:*",
        "work_type": "📋 Qanday ish kerak?",
        "mustaqil": "📝 Mustaqil ish", "kurs": "📚 Kurs ishi",
        "university": "🏛 *Universitetingiz* nomini kiriting:",
        "faculty": "🏫 *Fakultetingiz* nomini kiriting:",
        "direction": "📖 *Yo'nalishingizni* kiriting:",
        "subject": "📚 *Fan nomini* kiriting:",
        "student": "👤 *Ism va familiyangizni* kiriting:",
        "course": "🎓 *Nechanchi kursda* o'qiysiz? (1-4):",
        "study_type": "📅 O'qish shaklini tanlang:",
        "kunduzgi": "☀️ Kunduzgi", "sirtqi": "🌙 Sirtqi", "kechki": "🌆 Kechki",
        "topic": "✏️ *Ish mavzusini* kiriting:\n\n⚡ Mavzu qanchalik aniq bo'lsa, ish shunchalik sifatli bo'ladi!",
        "planning": "📋 *Reja tuzilmoqda...*\n\nBir daqiqa kuting ⏳",
        "gen_kirish": "✍️ *Kirish yozilmoqda...* (1/{total})",
        "gen_bolim": "✍️ *{nom} yozilmoqda...* ({step}/{total})",
        "gen_xulosa": "✍️ *Xulosa yozilmoqda...* ({step}/{total})",
        "gen_gloss": "✍️ *Glossariy tuzilmoqda...* ({step}/{total})",
        "gen_adab": "✍️ *Adabiyotlar ro'yxati...* ({step}/{total})",
        "combining": "📎 *Word fayl tayyorlanmoqda...*",
        "done": "✅ *Ish tayyor!*",
        "error": "❌ Xatolik yuz berdi. /start bilan qaytadan urinib ko'ring.",
        "invalid_course": "⚠️ 1 dan 4 gacha raqam kiriting.",
        "restart": "\n\n🔄 Yangi ish uchun: /start",
    },
    "ru": {
        "welcome": "👋 *Здравствуйте!*\n\nЯ бот для написания учебных работ по стандартам узбекских вузов.\n\n📝 *Самостоятельная работа* — 15–20 страниц\n📚 *Курсовая работа* — 25–35 страниц\n\n🌐 *Выберите язык:*",
        "work_type": "📋 Какая работа нужна?",
        "mustaqil": "📝 Самостоятельная работа", "kurs": "📚 Курсовая работа",
        "university": "🏛 Введите *название университета*:",
        "faculty": "🏫 Введите *название факультета*:",
        "direction": "📖 Введите *направление обучения*:",
        "subject": "📚 Введите *название предмета*:",
        "student": "👤 Введите *ваше имя и фамилию*:",
        "course": "🎓 На *каком курсе* вы учитесь? (1-4):",
        "study_type": "📅 Выберите форму обучения:",
        "kunduzgi": "☀️ Дневное", "sirtqi": "🌙 Заочное", "kechki": "🌆 Вечернее",
        "topic": "✏️ Введите *тему работы*:\n\n⚡ Чем точнее тема, тем качественнее работа!",
        "planning": "📋 *Составляется план...*\n\nОдну минуту ⏳",
        "gen_kirish": "✍️ *Пишется введение...* (1/{total})",
        "gen_bolim": "✍️ *Пишется {nom}...* ({step}/{total})",
        "gen_xulosa": "✍️ *Пишется заключение...* ({step}/{total})",
        "gen_gloss": "✍️ *Составляется глоссарий...* ({step}/{total})",
        "gen_adab": "✍️ *Список литературы...* ({step}/{total})",
        "combining": "📎 *Файл готовится...*",
        "done": "✅ *Работа готова!*",
        "error": "❌ Произошла ошибка. Попробуйте /start снова.",
        "invalid_course": "⚠️ Введите цифру от 1 до 4.",
        "restart": "\n\n🔄 Для новой работы: /start",
    },
    "en": {
        "welcome": "👋 *Hello!*\n\nI'm a bot for writing academic papers per Uzbekistan HEI standards.\n\n📝 *Independent work* — 15–20 pages\n📚 *Course work* — 25–35 pages\n\n🌐 *Choose language:*",
        "work_type": "📋 What type of work?",
        "mustaqil": "📝 Independent work", "kurs": "📚 Course work",
        "university": "🏛 Enter the *full university name*:",
        "faculty": "🏫 Enter your *faculty name*:",
        "direction": "📖 Enter your *study direction*:",
        "subject": "📚 Enter the *subject name*:",
        "student": "👤 Enter your *full name*:",
        "course": "🎓 Which *year* are you in? (1-4):",
        "study_type": "📅 Select study type:",
        "kunduzgi": "☀️ Full-time", "sirtqi": "🌙 Part-time", "kechki": "🌆 Evening",
        "topic": "✏️ Enter the *work topic*:\n\n⚡ The more specific, the better!",
        "planning": "📋 *Creating outline...*\n\nOne moment ⏳",
        "gen_kirish": "✍️ *Writing introduction...* (1/{total})",
        "gen_bolim": "✍️ *Writing {nom}...* ({step}/{total})",
        "gen_xulosa": "✍️ *Writing conclusion...* ({step}/{total})",
        "gen_gloss": "✍️ *Creating glossary...* ({step}/{total})",
        "gen_adab": "✍️ *References list...* ({step}/{total})",
        "combining": "📎 *Preparing file...*",
        "done": "✅ *Work is ready!*",
        "error": "❌ An error occurred. Try /start again.",
        "invalid_course": "⚠️ Please enter a number from 1 to 4.",
        "restart": "\n\n🔄 For a new work: /start",
    },
}

STUDY_TYPES = {
    "uz": {"kunduzgi": "kunduzgi", "sirtqi": "sirtqi", "kechki": "kechki"},
    "ru": {"kunduzgi": "дневное", "sirtqi": "заочное", "kechki": "вечернее"},
    "en": {"kunduzgi": "full-time", "sirtqi": "part-time", "kechki": "evening"},
}

def lang_kb():
    return InlineKeyboardMarkup([[
        InlineKeyboardButton("🇺🇿 O'zbek", callback_data="lang_uz"),
        InlineKeyboardButton("🇷🇺 Русский", callback_data="lang_ru"),
        InlineKeyboardButton("🇬🇧 English", callback_data="lang_en"),
    ]])

def work_kb(lang):
    return InlineKeyboardMarkup([
        [InlineKeyboardButton(T[lang]["mustaqil"], callback_data="work_mustaqil")],
        [InlineKeyboardButton(T[lang]["kurs"], callback_data="work_kurs")],
    ])

def study_kb(lang):
    return InlineKeyboardMarkup([[
        InlineKeyboardButton(T[lang]["kunduzgi"], callback_data="study_kunduzgi"),
        InlineKeyboardButton(T[lang]["sirtqi"],   callback_data="study_sirtqi"),
        InlineKeyboardButton(T[lang]["kechki"],   callback_data="study_kechki"),
    ]])

# ── Gemini ────────────────────────────────────────────────────────
def ask(prompt: str) -> str:
    return gemini.generate_content(prompt).text.strip()

# ── REJA ─────────────────────────────────────────────────────────
def plan_prompt(d: dict) -> str:
    s, t, lang, wt = d["subject"], d["topic"], d["lang"], d["work_type"]
    if lang == "uz":
        if wt == "mustaqil":
            return f'''"{s}" fanidan "{t}" mavzusida mustaqil ish rejasin FAQAT JSON formatida ber:
{{"b1": "Birinchi bo\'lim nomi", "b2": "Ikkinchi bo\'lim nomi", "b3": "Uchinchi bo\'lim nomi"}}'''
        else:
            return f'''"{s}" fanidan "{t}" mavzusida kurs ishi rejasin FAQAT JSON formatida ber:
{{"bob1": "I Bob nomi (nazariy)", "p11": "1.1 § nomi", "p12": "1.2 § nomi", "p13": "1.3 § nomi",
"bob2": "II Bob nomi (amaliy)", "p21": "2.1 § nomi", "p22": "2.2 § nomi", "p23": "2.3 § nomi"}}'''
    elif lang == "ru":
        if wt == "mustaqil":
            return f'''Для работы по предмету "{s}" на тему "{t}" дай план ТОЛЬКО в JSON:
{{"b1": "Название первого раздела", "b2": "Название второго раздела", "b3": "Название третьего раздела"}}'''
        else:
            return f'''Для курсовой по "{s}" на тему "{t}" дай план ТОЛЬКО в JSON:
{{"bob1": "Название I главы (теория)", "p11": "1.1 § название", "p12": "1.2 § название", "p13": "1.3 § название",
"bob2": "Название II главы (практика)", "p21": "2.1 § название", "p22": "2.2 § название", "p23": "2.3 § название"}}'''
    else:
        if wt == "mustaqil":
            return f'''For "{s}", topic "{t}", give outline ONLY in JSON:
{{"b1": "First section title", "b2": "Second section title", "b3": "Third section title"}}'''
        else:
            return f'''For course work on "{s}", topic "{t}", give outline ONLY in JSON:
{{"bob1": "Chapter I title (theory)", "p11": "1.1 § title", "p12": "1.2 § title", "p13": "1.3 § title",
"bob2": "Chapter II title (practical)", "p21": "2.1 § title", "p22": "2.2 § title", "p23": "2.3 § title"}}'''

# ── BO'LIM PROMPTLARI ─────────────────────────────────────────────
STYLE = {
    "uz": """MUHIM USLUB TALABLARI:
- Faqat uzun, oqimli paragraflar yoz. Har bir paragraf kamida 6-8 jumladan iborat bo'lsin.
- Ro'yxat (bullet, raqam) MUTLAQO ishlatma.
- Har bir paragraf oldingi paragrafdan mantiqiy davom etsin.
- Ilmiy, rasmiy o'zbek tilida yoz.
- Sarlavha yozma, faqat matn.
- Kamida 700 so'z.""",
    "ru": """ТРЕБОВАНИЯ К СТИЛЮ:
- Только длинные, связные абзацы. Каждый абзац минимум 6-8 предложений.
- Списки (маркированные, нумерованные) НЕ использовать.
- Каждый абзац логически продолжает предыдущий.
- Научный, официальный стиль на русском языке.
- Заголовок не писать, только текст.
- Минимум 700 слов.""",
    "en": """STYLE REQUIREMENTS:
- Only long, flowing paragraphs. Each paragraph minimum 6-8 sentences.
- NO lists (bullet points, numbered lists).
- Each paragraph logically continues from the previous.
- Academic, formal English style.
- No heading, only text.
- Minimum 700 words.""",
}

def write_kirish(d: dict) -> str:
    s, t, lang = d["subject"], d["topic"], d["lang"]
    if lang == "uz":
        p = f'''"{s}" fanidan "{t}" mavzusida mustaqil/kurs ishi uchun KIRISH bo\'limini yoz.
Quyidagilarni qamrab ol: mavzuning dolzarbligi va zamonaviy ahamiyati, mavzuning ilmiy o\'rganilganlik darajasi, ishning maqsadi, vazifalari, ob\'ekti va predmeti, nazariy va amaliy ahamiyati.
{STYLE[lang]}'''
    elif lang == "ru":
        p = f'''Напиши ВВЕДЕНИЕ для работы по предмету "{s}" на тему "{t}".
Включи: актуальность и современное значение темы, степень научной изученности, цель работы, задачи, объект и предмет, теоретическое и практическое значение.
{STYLE[lang]}'''
    else:
        p = f'''Write the INTRODUCTION for work on "{s}", topic "{t}".
Include: relevance and modern significance, degree of scientific study, goal, tasks, object and subject, theoretical and practical significance.
{STYLE[lang]}'''
    return ask(p)

def write_bolim(d: dict, bolim_nomi: str, context: str = "") -> str:
    s, t, lang = d["subject"], d["topic"], d["lang"]
    ctx = f"\n\nBu bo'lim avvalgi bo'limlar bilan bog'liq: {context}" if context and lang == "uz" else \
          f"\n\nЭтот раздел связан с предыдущими: {context}" if context and lang == "ru" else \
          f"\n\nThis section connects to previous sections: {context}" if context else ""
    if lang == "uz":
        p = f'''"{s}" fanidan "{t}" mavzusidagi yozma ish uchun "{bolim_nomi}" bo\'limini yoz.
Bu bo\'limda mavzuning ushbu jihatini chuqur ilmiy tahlil qil. Nazariy asoslar, ilmiy qarashlar, misollar va dalillar bilan boyit.{ctx}
{STYLE[lang]}'''
    elif lang == "ru":
        p = f'''Для работы по предмету "{s}" на тему "{t}" напиши раздел "{bolim_nomi}".
Дай глубокий научный анализ этого аспекта темы. Обогати теоретическими основами, научными взглядами, примерами и доказательствами.{ctx}
{STYLE[lang]}'''
    else:
        p = f'''For work on "{s}", topic "{t}", write the section "{bolim_nomi}".
Give deep scientific analysis of this aspect. Enrich with theoretical foundations, scholarly views, examples and evidence.{ctx}
{STYLE[lang]}'''
    return ask(p)

def write_bob_xulosa(d: dict, bob_nomi: str, parlar: list) -> str:
    s, t, lang = d["subject"], d["topic"], d["lang"]
    pars = ", ".join(parlar)
    if lang == "uz":
        p = f'''"{t}" mavzusidagi "{bob_nomi}" bobi bo\'yicha XULOSALAR yoz.
Bu bobda ko\'rib chiqilgan quyidagi mavzularning asosiy xulosalarini bir-biriga bog\'liq holda yoz: {pars}.
{STYLE[lang]}
Kamida 400 so\'z.'''
    elif lang == "ru":
        p = f'''Напиши ВЫВОДЫ по главе "{bob_nomi}" работы на тему "{t}".
Связно изложи основные выводы по следующим темам главы: {pars}.
{STYLE[lang]}
Минимум 400 слов.'''
    else:
        p = f'''Write CONCLUSIONS for chapter "{bob_nomi}" of the work on "{t}".
Coherently present main conclusions on the following topics: {pars}.
{STYLE[lang]}
Minimum 400 words.'''
    return ask(p)

def write_xulosa(d: dict, plan: dict) -> str:
    s, t, lang, wt = d["subject"], d["topic"], d["lang"], d["work_type"]
    if lang == "uz":
        p = f'''"{s}" fanidan "{t}" mavzusidagi yozma ishning umumiy XULOSA bo\'limini yoz.
Butun ish davomida ko\'rib chiqilgan asosiy fikrlar, ilmiy topilmalar va xulosalarni bir-biriga bog\'liq holda yoz. Amaliy tavsiyalar ham ber.
{STYLE[lang]}
Kamida 500 so\'z.'''
    elif lang == "ru":
        p = f'''Напиши общее ЗАКЛЮЧЕНИЕ работы по предмету "{s}" на тему "{t}".
Связно изложи основные идеи, научные выводы и результаты всей работы. Дай практические рекомендации.
{STYLE[lang]}
Минимум 500 слов.'''
    else:
        p = f'''Write the general CONCLUSION for work on "{s}", topic "{t}".
Coherently present the main ideas, scientific findings and results. Give practical recommendations.
{STYLE[lang]}
Minimum 500 words.'''
    return ask(p)

def write_glossariy(d: dict) -> str:
    s, t, lang = d["subject"], d["topic"], d["lang"]
    if lang == "uz":
        p = f'''"{s}" fanidan "{t}" mavzusiga oid 15-20 ta muhim atama uchun GLOSSARIY tuz.
Har bir atamani quyidagi formatda yoz (faqat shu format, boshqa narsa yo\'q):
Atama — Ta\'rif (2-3 jumlada ilmiy, aniq izoh).
Ro\'yxatni to\'g\'ridan to\'g\'ri boshla, kirish so\'z yozma.'''
    elif lang == "ru":
        p = f'''Составь ГЛОССАРИЙ из 15-20 важных терминов по предмету "{s}", тема "{t}".
Каждый термин в формате (только этот формат):
Термин — Определение (2-3 предложения, научное, точное).
Начни прямо со списка, без вводных слов.'''
    else:
        p = f'''Create a GLOSSARY of 15-20 important terms for "{s}", topic "{t}".
Each term in format (only this format):
Term — Definition (2-3 sentences, scientific, precise).
Start directly with the list, no introduction.'''
    return ask(p)

def write_adabiyotlar(d: dict) -> str:
    s, t, lang = d["subject"], d["topic"], d["lang"]
    if lang == "uz":
        p = f'''"{s}" fanidan "{t}" mavzusiga oid FOYDALANILGAN ADABIYOTLAR ro\'yxatini tuz.
Kamida 12 ta manba. Faqat quyidagi formatda:
1. Muallif I.O. Kitob nomi. – Shahar: Nashriyot, Yil.
Oxirida 4-5 ta internet manba:
13. Muallif. Maqola nomi // Sayt nomi. URL: https://... (Murojaat sanasi: kun.oy.yil)
Faqat ro\'yxat, boshqa narsa yozma.'''
    elif lang == "ru":
        p = f'''Составь СПИСОК ЛИТЕРАТУРЫ по предмету "{s}", тема "{t}".
Минимум 12 источников. Только в формате:
1. Автор И.О. Название. – Город: Издательство, Год.
В конце 4-5 интернет-источника. Только список.'''
    else:
        p = f'''Create REFERENCES list for "{s}", topic "{t}".
Minimum 12 sources. Only in format:
1. Author. Title. – City: Publisher, Year.
Add 4-5 internet sources at the end. Only the list.'''
    return ask(p)

# ── Word hujjat ───────────────────────────────────────────────────
def setup_doc(doc):
    s = doc.sections[0]
    s.left_margin = Cm(3); s.right_margin = Cm(1.5)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"; n.font.size = Pt(14)
    pf = n.paragraph_format
    pf.line_spacing = Pt(21); pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0); pf.space_after = Pt(0)

def add_pagenum(doc):
    for sec in doc.sections:
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.font.name = "Times New Roman"; r.font.size = Pt(12)
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
        ins = OxmlElement("w:instrText"); ins.text = "PAGE"
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
        r._r.extend([b, ins, e])

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(21)
    r = p.add_run(text.upper())
    r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(16)

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(21)
    r = p.add_run(text)
    r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(14)

def body(doc, text):
    for chunk in text.split("\n\n"):
        chunk = re.sub(r"^\s*[-•*]\s*", "", chunk).strip()
        chunk = re.sub(r"^\d+\.\s+", "", chunk).strip()
        chunk = re.sub(r"\*\*(.+?)\*\*", r"\1", chunk)
        if not chunk: continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = Pt(21)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(chunk)
        r.font.name = "Times New Roman"; r.font.size = Pt(14)

def ref_body(doc, text):
    for line in text.split("\n"):
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line).strip()
        if not line: continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = Pt(21)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Times New Roman"; r.font.size = Pt(14)

def muqova(doc, d):
    import datetime
    year = datetime.datetime.now().year
    lang = d["lang"]

    def cp(txt, bold=False, size=13, space=0):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(space)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(19)
        r = p.add_run(txt); r.bold = bold
        r.font.name = "Times New Roman"; r.font.size = Pt(size)

    def rp(txt, size=13):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.right_indent = Cm(1.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(19)
        r = p.add_run(txt)
        r.font.name = "Times New Roman"; r.font.size = Pt(size)

    if lang == "uz":
        cp("O'ZBEKISTON RESPUBLIKASI OLIY VA O'RTA MAXSUS TA'LIM VAZIRLIGI", bold=True, size=12)
    elif lang == "ru":
        cp("МИНИСТЕРСТВО ВЫСШЕГО И СРЕДНЕГО СПЕЦИАЛЬНОГО ОБРАЗОВАНИЯ\nРЕСПУБЛИКИ УЗБЕКИСТАН", bold=True, size=12)
    else:
        cp("MINISTRY OF HIGHER AND SECONDARY SPECIALISED EDUCATION\nOF THE REPUBLIC OF UZBEKISTAN", bold=True, size=12)

    cp(""); cp(d["university"].upper(), bold=True, size=13, space=6)
    cp(""); cp(d["faculty"], size=12); cp(d["direction"], size=12)
    cp(""); cp(""); cp("")

    if lang == "uz": cp(f'"{d["subject"]}"', size=12); cp("fanidan", size=12)
    elif lang == "ru": cp(f'По предмету: "{d["subject"]}"', size=12)
    else: cp(f'Subject: "{d["subject"]}"', size=12)

    cp(""); cp(f'"{d["topic"].upper()}"', bold=True, size=14, space=4); cp("")

    if lang == "uz": wname = "MUSTAQIL ISH" if d["work_type"] == "mustaqil" else "KURS ISHI"
    elif lang == "ru": wname = "САМОСТОЯТЕЛЬНАЯ РАБОТА" if d["work_type"] == "mustaqil" else "КУРСОВАЯ РАБОТА"
    else: wname = "INDEPENDENT WORK" if d["work_type"] == "mustaqil" else "COURSE WORK"
    cp(wname, bold=True, size=16, space=6)

    cp(""); cp(""); cp("")
    if lang == "uz": rp(f"{d['course']}-kurs {d['study_type']} talabasi:"); rp(d["student_name"], size=14)
    elif lang == "ru": rp(f"Студент {d['course']}-го курса ({d['study_type']}):"); rp(d["student_name"], size=14)
    else: rp(f"{d['course']}-year student ({d['study_type']}):"); rp(d["student_name"], size=14)

    if d["work_type"] == "kurs":
        cp("")
        if lang == "uz": rp("Ilmiy rahbar:")
        elif lang == "ru": rp("Научный руководитель:")
        else: rp("Scientific supervisor:")
        rp("______________________")

    cp(""); cp(""); cp(""); cp("")
    if lang == "uz": cp(f"Toshkent — {year}", size=13)
    elif lang == "ru": cp(f"Ташкент — {year}", size=13)
    else: cp(f"Tashkent — {year}", size=13)

def build_doc(d: dict, secs: dict, plan: dict) -> bytes:
    doc = Document()
    setup_doc(doc); add_pagenum(doc)
    lang = d["lang"]

    muqova(doc, d)

    # Kirish
    doc.add_page_break()
    if lang == "uz": h1(doc, "KIRISH")
    elif lang == "ru": h1(doc, "ВВЕДЕНИЕ")
    else: h1(doc, "INTRODUCTION")
    body(doc, secs["kirish"])

    if d["work_type"] == "mustaqil":
        for i in range(1, 4):
            doc.add_page_break()
            nom = plan.get(f"b{i}", f"Bo'lim {i}")
            h1(doc, f"{i}. {nom}")
            body(doc, secs[f"b{i}"])

        doc.add_page_break()
        if lang == "uz": h1(doc, "XULOSA")
        elif lang == "ru": h1(doc, "ЗАКЛЮЧЕНИЕ")
        else: h1(doc, "CONCLUSION")
        body(doc, secs["xulosa"])

        doc.add_page_break()
        if lang == "uz": h1(doc, "FOYDALANILGAN ADABIYOTLAR")
        elif lang == "ru": h1(doc, "СПИСОК ЛИТЕРАТУРЫ")
        else: h1(doc, "REFERENCES")
        ref_body(doc, secs["adabiyotlar"])

    else:  # kurs ishi
        doc.add_page_break()
        h1(doc, f"I BOB. {plan.get('bob1', '')}")
        for i in range(1, 4):
            h2(doc, f"1.{i} § {plan.get(f'p1{i}', '')}")
            body(doc, secs[f"p1{i}"])

        doc.add_page_break()
        if lang == "uz": h1(doc, "I BOB BO'YICHA XULOSALAR")
        elif lang == "ru": h1(doc, "ВЫВОДЫ ПО I ГЛАВЕ")
        else: h1(doc, "CHAPTER I CONCLUSIONS")
        body(doc, secs["bob1_xulosa"])

        doc.add_page_break()
        h1(doc, f"II BOB. {plan.get('bob2', '')}")
        for i in range(1, 4):
            h2(doc, f"2.{i} § {plan.get(f'p2{i}', '')}")
            body(doc, secs[f"p2{i}"])

        doc.add_page_break()
        if lang == "uz": h1(doc, "II BOB BO'YICHA XULOSALAR")
        elif lang == "ru": h1(doc, "ВЫВОДЫ ПО II ГЛАВЕ")
        else: h1(doc, "CHAPTER II CONCLUSIONS")
        body(doc, secs["bob2_xulosa"])

        doc.add_page_break()
        if lang == "uz": h1(doc, "XULOSA")
        elif lang == "ru": h1(doc, "ЗАКЛЮЧЕНИЕ")
        else: h1(doc, "CONCLUSION")
        body(doc, secs["xulosa"])

        doc.add_page_break()
        if lang == "uz": h1(doc, "GLOSSARIY")
        elif lang == "ru": h1(doc, "ГЛОССАРИЙ")
        else: h1(doc, "GLOSSARY")
        ref_body(doc, secs["glossariy"])

        doc.add_page_break()
        if lang == "uz": h1(doc, "FOYDALANILGAN ADABIYOTLAR")
        elif lang == "ru": h1(doc, "СПИСОК ЛИТЕРАТУРЫ")
        else: h1(doc, "REFERENCES")
        ref_body(doc, secs["adabiyotlar"])

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.read()

# ── Handlers ──────────────────────────────────────────────────────
async def cmd_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    ctx.user_data.clear()
    await update.message.reply_text(T["uz"]["welcome"], reply_markup=lang_kb(), parse_mode="Markdown")
    return LANG

async def cb_lang(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query; await q.answer()
    lang = q.data.replace("lang_", "")
    ctx.user_data["lang"] = lang
    await q.edit_message_text(T[lang]["work_type"], reply_markup=work_kb(lang), parse_mode="Markdown")
    return WORK_TYPE

async def cb_work(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query; await q.answer()
    lang = ctx.user_data["lang"]
    ctx.user_data["work_type"] = q.data.replace("work_", "")
    await q.edit_message_text(T[lang]["university"], parse_mode="Markdown")
    return UNIVERSITY

async def msg_university(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    lang = ctx.user_data["lang"]
    ctx.user_data["university"] = update.message.text.strip()
    await update.message.reply_text(T[lang]["faculty"], parse_mode="Markdown")
    return FACULTY

async def msg_faculty(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    lang = ctx.user_data["lang"]
    ctx.user_data["faculty"] = update.message.text.strip()
    await update.message.reply_text(T[lang]["direction"], parse_mode="Markdown")
    return DIRECTION

async def msg_direction(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    lang = ctx.user_data["lang"]
    ctx.user_data["direction"] = update.message.text.strip()
    await update.message.reply_text(T[lang]["subject"], parse_mode="Markdown")
    return SUBJECT

async def msg_subject(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    lang = ctx.user_data["lang"]
    ctx.user_data["subject"] = update.message.text.strip()
    await update.message.reply_text(T[lang]["student"], parse_mode="Markdown")
    return STUDENT_NAME

async def msg_student(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    lang = ctx.user_data["lang"]
    ctx.user_data["student_name"] = update.message.text.strip()
    await update.message.reply_text(T[lang]["course"], parse_mode="Markdown")
    return COURSE

async def msg_course(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    lang = ctx.user_data["lang"]
    text = update.message.text.strip()
    if not text.isdigit() or int(text) not in range(1, 5):
        await update.message.reply_text(T[lang]["invalid_course"])
        return COURSE
    ctx.user_data["course"] = text
    await update.message.reply_text(T[lang]["study_type"], reply_markup=study_kb(lang))
    return STUDY_TYPE

async def cb_study(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query; await q.answer()
    lang = ctx.user_data["lang"]
    ctx.user_data["study_type"] = STUDY_TYPES[lang][q.data.replace("study_", "")]
    await q.edit_message_text(T[lang]["topic"], parse_mode="Markdown")
    return TOPIC

async def msg_topic(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    lang = ctx.user_data["lang"]
    ctx.user_data["topic"] = update.message.text.strip()
    d = ctx.user_data
    wt = d["work_type"]

    status = await update.message.reply_text(T[lang]["planning"], parse_mode="Markdown")

    try:
        # 1. REJA
        logger.info("Reja tuzilmoqda...")
        raw = ask(plan_prompt(d))
        clean = re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()
        plan = json.loads(clean)
        d["plan"] = plan
        logger.info(f"Reja: {plan}")

        secs = {}
        total = 6 if wt == "mustaqil" else 12

        # 2. KIRISH
        await status.edit_text(T[lang]["gen_kirish"].format(total=total), parse_mode="Markdown")
        secs["kirish"] = write_kirish(d)
        logger.info(f"Kirish: {len(secs['kirish'])} belgi")

        if wt == "mustaqil":
            # 3 bo'lim
            prev = ""
            for i in range(1, 4):
                nom = plan.get(f"b{i}", f"Bo'lim {i}")
                await status.edit_text(T[lang]["gen_bolim"].format(nom=nom, step=i+1, total=total), parse_mode="Markdown")
                secs[f"b{i}"] = write_bolim(d, nom, prev)
                prev = nom
                logger.info(f"b{i}: {len(secs[f'b{i}'])} belgi")

            # Xulosa
            await status.edit_text(T[lang]["gen_xulosa"].format(step=5, total=total), parse_mode="Markdown")
            secs["xulosa"] = write_xulosa(d, plan)

            # Adabiyotlar
            await status.edit_text(T[lang]["gen_adab"].format(step=6, total=total), parse_mode="Markdown")
            secs["adabiyotlar"] = write_adabiyotlar(d)

        else:  # kurs ishi
            # I Bob paragraflar
            for i in range(1, 4):
                nom = plan.get(f"p1{i}", f"1.{i} §")
                await status.edit_text(T[lang]["gen_bolim"].format(nom=nom, step=i+1, total=total), parse_mode="Markdown")
                secs[f"p1{i}"] = write_bolim(d, nom)
                logger.info(f"p1{i}: {len(secs[f'p1{i}'])} belgi")

            # I Bob xulosa
            await status.edit_text(T[lang]["gen_bolim"].format(nom="I Bob xulosalari", step=5, total=total), parse_mode="Markdown")
            secs["bob1_xulosa"] = write_bob_xulosa(d, plan.get("bob1","I Bob"), [plan.get(f"p1{i}","") for i in range(1,4)])

            # II Bob paragraflar
            for i in range(1, 4):
                nom = plan.get(f"p2{i}", f"2.{i} §")
                await status.edit_text(T[lang]["gen_bolim"].format(nom=nom, step=i+5, total=total), parse_mode="Markdown")
                secs[f"p2{i}"] = write_bolim(d, nom)
                logger.info(f"p2{i}: {len(secs[f'p2{i}'])} belgi")

            # II Bob xulosa
            await status.edit_text(T[lang]["gen_bolim"].format(nom="II Bob xulosalari", step=9, total=total), parse_mode="Markdown")
            secs["bob2_xulosa"] = write_bob_xulosa(d, plan.get("bob2","II Bob"), [plan.get(f"p2{i}","") for i in range(1,4)])

            # Xulosa
            await status.edit_text(T[lang]["gen_xulosa"].format(step=10, total=total), parse_mode="Markdown")
            secs["xulosa"] = write_xulosa(d, plan)

            # Glossariy
            await status.edit_text(T[lang]["gen_gloss"].format(step=11, total=total), parse_mode="Markdown")
            secs["glossariy"] = write_glossariy(d)

            # Adabiyotlar
            await status.edit_text(T[lang]["gen_adab"].format(step=12, total=total), parse_mode="Markdown")
            secs["adabiyotlar"] = write_adabiyotlar(d)

        # 3. FAYL
        await status.edit_text(T[lang]["combining"], parse_mode="Markdown")
        docx_bytes = build_doc(d, secs, plan)
        prefix = "mustaqil_ish" if wt == "mustaqil" else "kurs_ishi"
        fname = f"{prefix}_{d['topic'][:20].replace(' ','_')}.docx"

        await status.edit_text(T[lang]["done"], parse_mode="Markdown")
        await update.message.reply_document(
            document=io.BytesIO(docx_bytes),
            filename=fname,
            caption=f"📄 {d['topic']}",
        )

    except Exception as e:
        logger.error(f"XATO: {type(e).__name__}: {e}")
        await status.edit_text(T[lang]["error"])

    await update.message.reply_text(T[lang]["restart"])
    return ConversationHandler.END

async def cmd_cancel(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    lang = ctx.user_data.get("lang", "uz")
    await update.message.reply_text(T[lang]["restart"])
    return ConversationHandler.END

def main():
    app = Application.builder().token(BOT_TOKEN).build()
    conv = ConversationHandler(
        entry_points=[CommandHandler("start", cmd_start)],
        states={
            LANG:         [CallbackQueryHandler(cb_lang,    pattern="^lang_")],
            WORK_TYPE:    [CallbackQueryHandler(cb_work,    pattern="^work_")],
            UNIVERSITY:   [MessageHandler(filters.TEXT & ~filters.COMMAND, msg_university)],
            FACULTY:      [MessageHandler(filters.TEXT & ~filters.COMMAND, msg_faculty)],
            DIRECTION:    [MessageHandler(filters.TEXT & ~filters.COMMAND, msg_direction)],
            SUBJECT:      [MessageHandler(filters.TEXT & ~filters.COMMAND, msg_subject)],
            STUDENT_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, msg_student)],
            COURSE:       [MessageHandler(filters.TEXT & ~filters.COMMAND, msg_course)],
            STUDY_TYPE:   [CallbackQueryHandler(cb_study,   pattern="^study_")],
            TOPIC:        [MessageHandler(filters.TEXT & ~filters.COMMAND, msg_topic)],
        },
        fallbacks=[CommandHandler("cancel", cmd_cancel)],
        allow_reentry=True,
    )
    app.add_handler(conv)
    logger.info("✅ Bot ishga tushdi!")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()
